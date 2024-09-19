import PyPDF2
from langdetect import detect
import openai
from openai import OpenAI
from docx import Document
import os
from dotenv import load_dotenv
load_dotenv()

client = OpenAI(
    api_key=os.getenv('api_key')
)

def extract_pdf_text(pdf_file):
    """Extracts text from the uploaded PDF."""
    print("extract_pdf_text called")
    reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text
        else:
            print(f"Warning: No text extracted from page {reader.pages.index(page)}")
    if not text:
        print("Warning: No text extracted from the PDF")
    return text


def generate_article_from_content(content, topic, language):
    """Generates an article based on the content and topic using OpenAI API."""
    prompt = f"Write an article titled '{topic}' focused on {topic}. Use the document content for reference but avoid mentioning irrelevant terms. Here's the reference content:\n{content} in {language}"

    
    try:
        completion = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an assistant that helps write articles."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=3000,
            temperature=0.7
        )


        message = completion.choices[0].message.content
        print(message )



        article = message
    except Exception as e:
        print(f"Error generating article: {e}")
        article = "Failed to generate article."
    
    return article

def detect_language(content):
    """Detect the language of the document content."""
    return detect(content)


def save_article_to_docx(article, topic):
    """Saves the generated article to a .docx file."""
    # Create a new Document object
    doc = Document()
    
    # Add title and article content to the document
    doc.add_heading(topic, level=1)
    doc.add_paragraph(article)

    # Define file path (ensure the directory exists)
    docx_file_path = os.path.join('/Users/mvp/Desktop/Article-AI/docx', f"{topic.replace(' ', '_')}_article.docx")
    
    # Save the document to the file path
    doc.save(docx_file_path)

    return docx_file_path
